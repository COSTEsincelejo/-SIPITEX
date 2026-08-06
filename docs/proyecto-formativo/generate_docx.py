#!/usr/bin/env python3
"""Genera PROYECTO_FORMATIVO_SIPITEX_PUNTOS_1_9.docx a partir del contenido SIPITEX."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[2]
DIAG = ROOT / "docs" / "diagramas"
OUT = Path(__file__).resolve().parent / "PROYECTO_FORMATIVO_SIPITEX_PUNTOS_1_9.docx"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_custom(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=16 if level == 1 else 13, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))
    return p


def add_para(doc, text, bold=False, size=11, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            set_run_font(run)


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=10, bold=True)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Cm(w)
    doc.add_paragraph()
    return table


def add_image(doc, path: Path, width_inches=6.2, caption=None):
    if not path.exists():
        add_para(doc, f"[Imagen no encontrada: {path.name}]", bold=True)
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    if caption:
        cap = doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cap.add_run(caption)
        set_run_font(r, size=9, bold=False, color=RGBColor(0x55, 0x55, 0x55))


def add_uc_block(doc, nombre, actores, funcion, descripcion, refs):
    add_heading_custom(doc, nombre, level=3)
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Nombre", nombre],
            ["Actores", actores],
            ["Función", funcion],
            ["Descripción", descripcion],
            ["Referencias", refs],
        ],
        col_widths=[3.5, 13],
    )


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    # Portada
    for _ in range(2):
        doc.add_paragraph()
    add_para(doc, "SERVICIO NACIONAL DE APRENDIZAJE — SENA", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "CMTC · Programa ADSO", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "PROYECTO FORMATIVO", bold=True, size=18, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(
        doc,
        "SIPITEX — Sistema Integrado de Producción e Inventario Textil",
        bold=True,
        size=16,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )
    doc.add_paragraph()
    add_para(doc, "Presentado por:", size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Cristian Camilo Baena Ruiz", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "cristianccbr@gmail.com", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph()
    add_para(doc, "Documento: puntos 1 al 9", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, "Versión 1.0 · Julio 2026", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()

    # Contenido
    add_heading_custom(doc, "Contenido")
    for item in [
        "1. Levantamiento de información",
        "2. Informe de requerimientos",
        "3. Hardware del cliente",
        "4. Diagrama de Gantt",
        "5. Casos de uso",
        "6. Diagrama de flujo",
        "7. Diagrama de clases",
        "8. Diagrama de distribución",
        "9. Modelo entidad relación",
    ]:
        add_para(doc, item)
    doc.add_page_break()

    # ========== 1 ==========
    add_heading_custom(doc, "1. Levantamiento de información")
    add_heading_custom(doc, "1.1 Contexto del problema", level=2)
    add_para(
        doc,
        "El Centro de Manufactura en Confección y Textiles (CMTC) del SENA requiere un sistema "
        "informático que permita controlar de forma integrada el inventario de materias primas "
        "textiles, las órdenes de producción, el cálculo de requerimientos (MRP/BOM), el registro "
        "de producción por fichas de aprendices, las inspecciones de calidad y los reportes KPI.",
    )
    add_para(doc, "Actualmente parte de esta información se maneja de forma dispersa, lo que genera:")
    add_bullets(
        doc,
        [
            "Desconocimiento del stock real en tiempo real.",
            "Dificultad para aprobar o rechazar salidas de material con trazabilidad.",
            "Falta de avance consolidado de órdenes frente a la meta.",
            "Reportes lentos o incompletos para la coordinación académica-productiva.",
        ],
    )

    add_heading_custom(doc, "1.2 Técnicas de recolección", level=2)
    add_table(
        doc,
        ["Técnica", "Descripción", "Resultado"],
        [
            [
                "Entrevistas",
                "Conversaciones con instructores, bodeguero y administración",
                "Roles, permisos y flujos diarios",
            ],
            [
                "Observación",
                "Proceso de solicitud de materiales y registro de producción",
                "Flujo solicitud → aprobación → descuento",
            ],
            [
                "Análisis documental",
                "Formatos de inventario, fichas y órdenes",
                "Atributos mínimos de entidades",
            ],
            [
                "Benchmark",
                "Referencia a sistemas de inventario previos",
                "Adaptación al dominio textil",
            ],
        ],
    )

    add_heading_custom(doc, "1.3 Preguntas guía", level=2)
    add_bullets(
        doc,
        [
            "¿Quién solicita materiales y quién los aprueba?",
            "¿Cómo se identifica un material (código, nombre, unidad, stock mínimo)?",
            "¿Qué información debe llevar una orden de producción?",
            "¿Cómo se relaciona una ficha de aprendices con una orden?",
            "¿Qué criterios de calidad se registran (aprobado / reproceso)?",
            "¿Qué reportes necesita la administración?",
            "¿El sistema debe operar en intranet sin internet externo?",
            "¿Qué roles existen y qué puede hacer cada uno?",
        ],
    )

    add_heading_custom(doc, "1.4 Hallazgos principales", level=2)
    add_table(
        doc,
        ["Hallazgo", "Impacto en el diseño"],
        [
            ["Tres roles: Administrador, Bodeguero, Instructor", "Cookie Auth + control por rol"],
            ["Stock crítico y estado físico", "Niveles Agotado / Por agotarse / Normal"],
            ["Salidas ligadas a una orden", "Entidad MaterialRequest"],
            ["Necesidad de BOM por producto", "Módulo MRP con BomItem"],
            ["Producción por ficha", "Ficha y ProductionSession"],
            ["Reportes filtrables", "QuestPDF / ClosedXML"],
            ["Despliegue reproducible", "Docker Compose (RNF07)"],
        ],
    )

    add_heading_custom(doc, "1.5 Conclusión", level=2)
    add_para(
        doc,
        "Se confirma la necesidad de un sistema web monolítico por capas (SIPITEX) orientado a "
        "intranet, con SQLite, autenticación por sesión y módulos de Inventario, Órdenes, MRP, "
        "Fichas, Calidad, Estadísticas, Reportes, Alertas y Usuarios.",
    )
    doc.add_page_break()

    # ========== 2 ==========
    add_heading_custom(doc, "2. Informe de requerimientos")
    add_para(doc, "SIPITEX — INFORME DE REQUERIMIENTOS", bold=True, size=12)
    add_para(doc, "Cristian Camilo Baena Ruiz · cristianccbr@gmail.com")

    add_heading_custom(doc, "Introducción", level=2)
    add_para(
        doc,
        "El uso de software para el control de producción e inventario textil facilitará las "
        "actividades de administración, bodega e instructores del CMTC. Los beneficiados "
        "principales necesitan precisión y oportunidad en la información de materiales, órdenes "
        "y producción. Con SIPITEX se cubre la necesidad de un sistema web integrado que el "
        "centro no posee actualmente.",
    )

    add_heading_custom(doc, "Propósito", level=2)
    add_bullets(
        doc,
        [
            "Control de materiales (stock, mínimos, estado físico).",
            "Gestión de órdenes de producción y su avance.",
            "Cálculo de requerimientos (MRP) a partir del BOM.",
            "Registro de producción por fichas e inspecciones de calidad.",
            "Consultas, reportes KPI y alertas.",
        ],
    )

    add_heading_custom(doc, "Ámbito del sistema", level=2)
    add_para(
        doc,
        "Dentro del alcance: usuarios/roles, inventario, solicitudes, órdenes, MRP/BOM, fichas, "
        "calidad, reportes, alertas y estadísticas. Fuera de alcance: facturación, nómina, ERP "
        "externo y aplicación móvil nativa.",
    )

    add_heading_custom(doc, "2.1 Perspectiva del producto", level=2)
    add_table(
        doc,
        ["Capa", "Proyecto", "Responsabilidad"],
        [
            ["Presentación", "Sipitex.Web", "Controllers, Razor Views, Cookie Auth"],
            ["Aplicación", "Sipitex.Application", "Servicios, DTOs, reglas de negocio"],
            ["Dominio", "Sipitex.Domain", "Entidades y enums"],
            ["Infraestructura", "Sipitex.Infrastructure", "EF Core, SQLite, repositorios"],
        ],
    )

    add_heading_custom(doc, "2.2 Funciones del sistema", level=2)
    add_bullets(
        doc,
        [
            "Autenticar usuarios y controlar acceso por rol/permiso.",
            "Registrar y consultar materiales.",
            "Crear órdenes de producción.",
            "Mantener BOM y simular MRP.",
            "Solicitar, aprobar o rechazar salidas de bodega.",
            "Registrar producción por ficha y avance de orden.",
            "Registrar inspecciones de calidad.",
            "Generar reportes filtrables y alertas.",
        ],
    )

    add_heading_custom(doc, "2.3 Características de los usuarios", level=2)
    add_para(
        doc,
        "Interfaces intuitivas y de alto grado de usabilidad. Objetivo de aprendizaje: menos de "
        "4 horas con capacitación básica.",
    )

    add_heading_custom(doc, "2.4 Restricciones", level=2)
    add_bullets(
        doc,
        [
            "Componentes sin licenciamiento comercial obligatorio.",
            "Modelo cliente/servidor sobre HTTP/HTTPS.",
            "SQLite para desarrollo e intranet.",
            "Autenticación por cookies.",
            "Despliegue opcional con Docker Compose.",
        ],
    )

    add_heading_custom(doc, "3.1 Requisitos funcionales", level=2)
    rf_rows = [
        ["RF01", "Usuarios", "CRUD usuarios por rol", "Alta"],
        ["RF02", "Usuarios", "Iniciar/cerrar sesión con cookies", "Alta"],
        ["RF03", "Usuarios", "Permisos extendidos desde administrador", "Alta"],
        ["RF04", "Inventario", "Registrar material", "Alta"],
        ["RF05", "Inventario", "Consultar stock y filtrar por nivel", "Alta"],
        ["RF06", "Inventario", "Actualizar estado físico", "Media"],
        ["RF07", "Inventario", "Ajustar stock y fecha de entrada", "Media"],
        ["RF08", "Salida", "Solicitar material para una orden", "Alta"],
        ["RF09", "Salida", "Aprobar/rechazar y descontar stock", "Alta"],
        ["RF10", "Órdenes", "Crear orden de producción", "Alta"],
        ["RF11", "Órdenes", "Registrar avance y finalizar meta", "Alta"],
        ["RF12", "MRP", "Mantener BOM por producto", "Alta"],
        ["RF13", "MRP", "Simular requerimiento neto", "Alta"],
        ["RF14", "Fichas", "Asociar ficha a proceso/instructor/orden", "Alta"],
        ["RF15", "Producción", "Registrar sesión diaria", "Alta"],
        ["RF16", "Calidad", "Inspección; reproceso con motivo", "Media"],
        ["RF17", "Reportes", "Exportar PDF/Excel", "Alta"],
        ["RF18", "Reportes", "Filtros por período/instructor/ficha", "Alta"],
        ["RF19", "Alertas", "Preferencias y evaluación de eventos", "Media"],
        ["RF20", "Estadísticas", "Dashboard KPI", "Alta"],
    ]
    add_table(doc, ["ID", "Módulo", "Descripción", "Prioridad"], rf_rows)

    add_heading_custom(doc, "3.2 Requisitos no funcionales", level=2)
    add_table(
        doc,
        ["ID", "Descripción"],
        [
            ["RNF01", "Respuesta percibida < 2 s en intranet"],
            ["RNF02", "Sesión con expiración de 8 horas"],
            ["RNF03", "Control de acceso por rol y permisos"],
            ["RNF04", "Acceso desde PCs de la intranet"],
            ["RNF05", "UI responsiva"],
            ["RNF06", "Código modular por capas"],
            ["RNF07", "Docker Compose"],
            ["RNF08", "Integridad con EF Core / SQLite"],
        ],
    )

    add_heading_custom(doc, "3.3 Interfaces externas", level=2)
    add_table(
        doc,
        ["Tipo", "Detalle"],
        [
            ["Usuario", "Navegador web (HTML, CSS, JavaScript)"],
            ["Hardware", "Servidor HTTP o contenedor Docker"],
            ["Software", ".NET 10, EF Core, QuestPDF, ClosedXML, MailKit"],
            ["Comunicaciones", "HTTP(S) local; SMTP opcional"],
        ],
    )

    add_heading_custom(doc, "3.4 Ciclo de vida", level=2)
    add_para(
        doc,
        "Metodología cascada: análisis → diseño → implementación → pruebas → despliegue, "
        "documentada en la carpeta docs/ del repositorio.",
    )
    doc.add_page_break()

    # ========== 3 ==========
    add_heading_custom(doc, "3. Hardware del cliente")
    add_para(
        doc,
        "Para la instalación y puesta en marcha de SIPITEX, el cliente (SENA CMTC) debe "
        "provisionar un servidor o estación con las siguientes características:",
    )
    add_table(
        doc,
        ["Componente", "Mínimo", "Recomendado"],
        [
            ["Procesador", "2 núcleos x86-64", "4 núcleos o superior"],
            ["Memoria RAM", "4 GB", "8 GB"],
            ["Almacenamiento", "20 GB libres (SSD)", "40 GB SSD"],
            ["Sistema operativo", "Windows 10/11 o Ubuntu 22.04+", "Windows Server / Ubuntu LTS"],
            ["Red", "Ethernet 100 Mbps LAN", "Gigabit Ethernet"],
            ["Runtime", ".NET 10 Runtime o Docker", ".NET 10 + Docker Compose"],
            ["Clientes", "Chrome / Edge / Firefox", "Misma especificación en bodega/aula"],
        ],
    )
    add_bullets(
        doc,
        [
            "La base de datos SQLite (sipitex.db) reside en el servidor; respaldo diario recomendado.",
            "En Docker, el volumen sipitex-data persiste la BD.",
            "No se requiere internet externo para el funcionamiento core (SMTP es opcional).",
        ],
    )
    doc.add_page_break()

    # ========== 4 ==========
    add_heading_custom(doc, "4. Diagrama de Gantt")
    add_heading_custom(doc, "4.1 Cronograma del proyecto", level=2)
    add_table(
        doc,
        ["Fase", "Actividades", "Duración"],
        [
            ["1. Requisitos", "Levantamiento, informe RF/RNF, aprobación", "4 semanas"],
            ["2. Diseño", "Arquitectura, ER, casos de uso, clases, flujos", "4 semanas"],
            ["3. Implementación", "Capas y módulos de negocio", "9–10 semanas"],
            ["4. Pruebas", "Unitarias, funcionales, correcciones", "4 semanas"],
            ["5. Despliegue", "Intranet/Docker, capacitación, entrega", "2 semanas"],
        ],
    )
    add_image(doc, DIAG / "11-gantt.png", 6.3, "Figura 4.1 — Diagrama de Gantt SIPITEX")

    add_heading_custom(doc, "4.2 Cuadro comparativo de tecnologías", level=2)
    add_table(
        doc,
        ["Software", "Definición", "Ventajas", "Desventajas"],
        [
            ["SQLite (elegido)", "Motor SQL embebido", "Bajo footprint, fácil respaldo", "Menor concurrencia"],
            ["PostgreSQL", "SGBD relacional libre", "Potente y escalable", "Administración extra"],
            ["MySQL", "SGBD open source", "Rápido en web", "Configuración adicional"],
            ["SQL Server", "SGBD Microsoft", "Integración Windows", "Costo de licencias"],
            ["Oracle", "SGBD empresarial", "Alta disponibilidad", "Precio excesivo para CMTC"],
            ["VS / VS Code", "IDE C# / .NET", "Productividad ASP.NET", "VS completo es pesado"],
            ["Windows 10/11", "SO cliente/servidor", "Familiar en el centro", "Licencia Microsoft"],
            ["Linux", "SO libre tipo Unix", "Gratuito y robusto", "Curva de aprendizaje"],
        ],
    )
    add_para(
        doc,
        "Decisión: ASP.NET Core MVC + EF Core + SQLite + Docker Compose, por equilibrio entre "
        "costo, facilidad de despliegue en intranet y alineación con ADSO.",
        bold=True,
    )
    doc.add_page_break()

    # ========== 5 ==========
    add_heading_custom(doc, "5. Casos de uso")
    add_heading_custom(doc, "5.1 Diagrama de casos de uso", level=2)
    add_image(doc, DIAG / "01-casos-de-uso.png", 6.0, "Figura 5.1 — Casos de uso SIPITEX")

    add_heading_custom(doc, "5.2 Matriz actor ↔ caso de uso", level=2)
    add_table(
        doc,
        ["Caso de uso", "Admin", "Bodeguero", "Instructor"],
        [
            ["1. Iniciar sesión", "✓", "✓", "✓"],
            ["2. Gestionar usuarios", "✓", "", ""],
            ["3. Registrar materiales", "✓", "✓", "permiso"],
            ["4. Consultar stock", "✓", "✓", "✓"],
            ["5. Solicitar material", "✓", "", "✓"],
            ["6. Aprobar / rechazar", "✓", "✓", "permiso"],
            ["7. Crear orden", "✓", "", ""],
            ["8. Registrar producción", "✓", "", "✓"],
            ["9. BOM / MRP", "✓", "✓", "permiso"],
            ["10. Control de calidad", "✓", "", "✓"],
            ["11. Descargar reportes", "✓", "✓", "✓"],
            ["12. Configurar alertas", "✓", "parcial", "parcial"],
        ],
    )

    add_heading_custom(doc, "5.3 Descripción de casos de uso", level=2)
    add_uc_block(
        doc,
        "CU-01 — Iniciar sesión",
        "Administrador, Bodeguero, Instructor",
        "Autenticar al usuario y abrir sesión",
        "El usuario ingresa correo y contraseña. El sistema valida credenciales, crea cookie de autenticación y redirige según rol.",
        "RF02",
    )
    add_uc_block(
        doc,
        "CU-02 — Gestión de usuarios",
        "Administrador",
        "Mantenimiento de usuarios y permisos",
        "Crear, editar, activar/desactivar usuarios; asignar rol y permisos extendidos.",
        "RF01, RF03",
    )
    add_uc_block(
        doc,
        "CU-03 — Administración de materiales",
        "Administrador, Bodeguero",
        "Mantener inventario de materias primas",
        "Registrar material (código, nombre, unidad, stock, mínimo, estado). Consultar y filtrar por nivel. Ajustar entradas y estado físico.",
        "RF04–RF07",
    )
    add_uc_block(
        doc,
        "CU-04 — Solicitar material",
        "Instructor, Administrador",
        "Pedir salida de bodega asociada a una orden",
        "Selecciona orden y material (o crea material), indica cantidad y envía solicitud en estado Pendiente.",
        "RF08",
    )
    add_uc_block(
        doc,
        "CU-05 — Aprobar / rechazar solicitud",
        "Bodeguero, Administrador",
        "Resolver solicitudes de material",
        "Al aprobar se descuenta stock; al rechazar se registra motivo. Queda trazabilidad por orden.",
        "RF09",
    )
    add_uc_block(
        doc,
        "CU-06 — Crear orden de producción",
        "Administrador",
        "Abrir orden con meta de unidades",
        "Se registra número de orden, producto, cantidad total, fecha límite y estado. El avance se actualiza con sesiones de producción.",
        "RF10, RF11",
    )
    add_uc_block(
        doc,
        "CU-07 — BOM / MRP",
        "Administrador, Bodeguero",
        "Definir materiales por producto y calcular faltantes",
        "Se agregan ítems al BOM. La simulación MRP compara requerimiento neto vs stock disponible.",
        "RF12, RF13",
    )
    add_uc_block(
        doc,
        "CU-08 — Registrar producción",
        "Instructor, Administrador",
        "Sesión diaria de producción",
        "Se asocia ficha a orden; se registran unidades y observaciones; se actualiza avance; si se alcanza la meta se finaliza la orden.",
        "RF14, RF15",
    )
    add_uc_block(
        doc,
        "CU-09 — Control de calidad",
        "Instructor, Administrador",
        "Registrar inspección",
        "Unidades inspeccionadas y resultado (Aprobado / Reproceso). En reproceso se exigen motivo y responsable.",
        "RF16",
    )
    add_uc_block(
        doc,
        "CU-10 — Reportes y alertas",
        "Todos (según módulo)",
        "Exportar información y recibir avisos",
        "Exportación PDF/Excel con filtros. Preferencias de alerta (stock bajo, solicitudes pendientes, órdenes atrasadas, reprocesos).",
        "RF17–RF20",
    )
    doc.add_page_break()

    # ========== 6 ==========
    add_heading_custom(doc, "6. Diagrama de flujo")
    add_heading_custom(doc, "6.1 Flujo principal", level=2)
    add_para(
        doc,
        "El siguiente diagrama describe el flujo desde el inicio de sesión hasta la resolución "
        "de una solicitud de material, el registro de producción y la inspección de calidad.",
    )
    add_image(doc, DIAG / "12-flujo-solicitud.png", 5.8, "Figura 6.1 — Diagrama de flujo SIPITEX")

    add_heading_custom(doc, "6.2 Flujos de secuencia", level=2)
    add_para(doc, "Complementan el flujo lógico con el orden de mensajes entre capas:")
    for name, file, caption in [
        ("Login", "04-secuencia-login.png", "Figura 6.2 — Secuencia login"),
        ("Solicitar material", "05-secuencia-solicitar.png", "Figura 6.3 — Solicitar material"),
        ("Aprobar solicitud", "06-secuencia-aprobar.png", "Figura 6.4 — Aprobar solicitud"),
        ("Crear orden", "07-secuencia-crear-orden.png", "Figura 6.5 — Crear orden"),
    ]:
        add_para(doc, name, bold=True)
        add_image(doc, DIAG / file, 5.8, caption)

    add_heading_custom(doc, "6.3 Descripción narrativa — solicitud", level=2)
    add_bullets(
        doc,
        [
            "El Instructor inicia sesión.",
            "Selecciona una orden de producción activa.",
            "Elige material existente o registra uno nuevo.",
            "Indica cantidad y envía la solicitud (estado Pendiente).",
            "El Bodeguero revisa la solicitud.",
            "Si aprueba, el sistema descuenta stock y marca Aprobada.",
            "Si rechaza, registra motivo y marca Rechazada.",
            "Queda trazabilidad asociada a la orden.",
        ],
    )
    doc.add_page_break()

    # ========== 7 ==========
    add_heading_custom(doc, "7. Diagrama de clases")
    add_heading_custom(doc, "7.1 Clases del dominio", level=2)
    add_image(doc, DIAG / "02-clases-dominio.png", 6.2, "Figura 7.1 — Diagrama de clases")

    add_heading_custom(doc, "7.2 Vista de capas de aplicación", level=2)
    add_image(doc, DIAG / "03-capas-aplicacion.png", 6.0, "Figura 7.2 — Capas de aplicación")

    add_heading_custom(doc, "7.3 Resumen de entidades", level=2)
    add_table(
        doc,
        ["Clase", "Responsabilidad"],
        [
            ["User", "Usuarios, rol, permisos, ficha asignada"],
            ["Material", "Inventario de materias primas"],
            ["BomItem", "Material por unidad de producto"],
            ["ProductionOrder", "Órdenes de producción y avance"],
            ["MaterialRequest", "Solicitudes de salida de bodega"],
            ["Ficha", "Proceso / grupo asociado a instructor y orden"],
            ["ProductionSession", "Registro diario de unidades"],
            ["QualityRecord", "Inspecciones de calidad"],
            ["AlertPreference / AlertDelivery", "Preferencias y envíos de alertas"],
        ],
    )
    add_heading_custom(doc, "7.4 Patrones aplicados", level=2)
    add_bullets(doc, ["Repository + Unit of Work", "DTO", "Dependency Injection"])
    doc.add_page_break()

    # ========== 8 ==========
    add_heading_custom(doc, "8. Diagrama de distribución")
    add_heading_custom(doc, "8.1 Despliegue físico / lógico", level=2)
    add_image(doc, DIAG / "13-distribucion.png", 6.0, "Figura 8.1 — Diagrama de distribución")

    add_heading_custom(doc, "8.2 Arquitectura por capas", level=2)
    add_image(doc, DIAG / "00-arquitectura.png", 5.5, "Figura 8.2 — Arquitectura por capas")

    add_heading_custom(doc, "8.3 Nodos de despliegue", level=2)
    add_table(
        doc,
        ["Nodo", "Componentes", "Protocolo"],
        [
            ["PCs de usuarios", "Navegador web", "HTTP/HTTPS"],
            ["Switch / LAN CMTC", "Red local", "Ethernet"],
            ["Servidor de aplicaciones", "Kestrel / Docker · Sipitex.Web", "Puerto 8080 o IIS"],
            ["Almacenamiento", "sipitex.db (SQLite) / volumen Docker", "Archivo local"],
            ["SMTP (opcional)", "MailKit", "SMTP"],
        ],
    )

    add_heading_custom(doc, "8.4 Opciones de publicación", level=2)
    add_bullets(
        doc,
        [
            "Local: dotnet run en src/Sipitex.Web",
            "Publish: dotnet publish -c Release + IIS o ejecutable",
            "Docker Compose: docker compose up --build → http://localhost:8080",
        ],
    )
    doc.add_page_break()

    # ========== 9 ==========
    add_heading_custom(doc, "9. Modelo entidad relación")
    add_heading_custom(doc, "9.1 Diagrama ER", level=2)
    add_image(doc, DIAG / "10-entidad-relacion.png", 6.2, "Figura 9.1 — Modelo entidad-relación")

    add_heading_custom(doc, "9.2 Relaciones principales", level=2)
    add_table(
        doc,
        ["Relación", "Cardinalidad", "Descripción"],
        [
            ["USERS → FICHAS", "0..1", "Ficha asignada al usuario"],
            ["MATERIALS → BOM_ITEMS", "1..*", "Material usado en BOM"],
            ["MATERIALS → MATERIAL_REQUESTS", "1..*", "Material pedido"],
            ["PRODUCTION_ORDERS → MATERIAL_REQUESTS", "1..*", "Solicitudes de la orden"],
            ["PRODUCTION_ORDERS → QUALITY_RECORDS", "1..*", "Inspecciones"],
            ["PRODUCTION_ORDERS → FICHAS", "1..*", "Fichas asignadas"],
            ["FICHAS → PRODUCTION_SESSIONS", "1..*", "Sesiones de producción"],
            ["USERS → ALERT_PREFERENCES", "1..*", "Preferencias de alerta"],
        ],
    )

    add_heading_custom(doc, "9.3 Diccionario resumido de tablas", level=2)
    add_table(
        doc,
        ["Tabla", "Clave", "Campos clave"],
        [
            ["USERS", "Id", "Nombre, Email, PasswordHash, Rol, PermisosExtendidos, IsActive"],
            ["MATERIALS", "Id", "Code, Name, Unit, Stock, MinStock, Status, LastEntryDate"],
            ["BOM_ITEMS", "Id", "ProductName, MaterialId, QuantityPerUnit, Unit"],
            ["PRODUCTION_ORDERS", "Id", "OrderNumber, ProductName, TotalQuantity, ProducedQuantity, Status"],
            ["MATERIAL_REQUESTS", "Id", "MaterialId, ProductionOrderId, Quantity, Status, CreatedAt"],
            ["FICHAS", "Id", "FichaCode, ProcessName, InstructorName, ProductionOrderId"],
            ["PRODUCTION_SESSIONS", "Id", "FichaId, ProductionOrderId, Units, Observations, SessionDate"],
            ["QUALITY_RECORDS", "Id", "ProductionOrderId, UnitsInspected, Result, MotivoReproceso"],
            ["ALERT_PREFERENCES", "Id", "UserId, AlertType, Enabled"],
            ["ALERT_DELIVERIES", "Id", "UserId, AlertType, Subject, Body, SentAt, Channel"],
        ],
    )

    add_heading_custom(doc, "9.4 Integridad", level=2)
    add_bullets(
        doc,
        [
            "Claves foráneas gestionadas por EF Core.",
            "Transacciones en operaciones críticas (aprobación + descuento de stock).",
            "Migraciones EF Core (MigrateAsync) al arrancar; baseline automático para BD legacy.",
        ],
    )
    doc.add_page_break()

    # Apéndices
    add_heading_custom(doc, "Apéndice A — Credenciales de demostración")
    add_table(
        doc,
        ["Rol", "Correo", "Contraseña"],
        [
            ["Administrador", "admin@sipitex.test", "Admin123!"],
            ["Instructor", "instructor@sipitex.test", "Instructor123!"],
            ["Bodeguero", "bodega@sipitex.test", "Bodega123!"],
        ],
    )

    add_heading_custom(doc, "Apéndice B — Cómo ejecutar")
    add_para(doc, "Ejecución local:", bold=True)
    add_para(doc, "cd src/Sipitex.Web && dotnet run")
    add_para(doc, "Docker Compose:", bold=True)
    add_para(doc, "docker compose up --build")
    add_para(doc, "Abrir http://localhost:8080")

    add_para(
        doc,
        "Fin del documento formativo SIPITEX — puntos 1 a 9 · v1.0",
        bold=True,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )

    doc.save(OUT)
    print(f"Generado: {OUT}")


if __name__ == "__main__":
    build()
