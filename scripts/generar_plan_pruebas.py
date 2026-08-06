#!/usr/bin/env python3
"""Genera el Plan de Pruebas de SIPITEX (DOCX + PDF) a partir de la plantilla SDS."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.platypus import (
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)

OUT_DIR = Path("/workspace/docs/entregables")
ARTIFACT_DIR = Path("/opt/cursor/artifacts")
OUT_DIR.mkdir(parents=True, exist_ok=True)
ARTIFACT_DIR.mkdir(parents=True, exist_ok=True)

TITLE = "PLAN DE PRUEBAS — SIPITEX"
SUBTITLE = "Sistema Integrado de Producción e Inventario Textil"


def set_run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def add_heading_doc(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))
    return p


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p


def add_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for run in p.runs:
                set_run_font(run, size=9, bold=True)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = str(val)
            for p in cells[c_idx].paragraphs:
                for run in p.runs:
                    set_run_font(run, size=9)
    doc.add_paragraph()
    return table


def build_docx(path: Path):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # Portada
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SENA — CMTC · Programa ADSO")
    set_run_font(r, size=12, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("SISTEMA INTEGRADO DE GESTIÓN\nCONTROL DOCUMENTAL")
    set_run_font(r, size=10, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("GUÍA PARA EL DOCUMENTO DE PLAN DE PRUEBAS DE SOFTWARE")
    set_run_font(r, size=11, bold=True)

    add_para(doc, "Código: SIP-TIC-GUI-004  |  Versión: 1.0")
    add_para(doc, "Elaborado con base en la plantilla SDS-TIC-GUI-004 / 114-GTI-GUI-08 y adaptado al sistema SIPITEX.")
    add_para(doc, "Fecha: Julio 2026")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(TITLE)
    set_run_font(r, size=18, bold=True, color=RGBColor(0x1A, 0x3A, 0x5C))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(SUBTITLE)
    set_run_font(r, size=12)

    doc.add_page_break()

    # 1. OBJETIVO
    add_heading_doc(doc, "1. OBJETIVO")
    add_para(
        doc,
        "Verificar el correcto funcionamiento del Sistema Integrado de Producción e Inventario Textil "
        "(SIPITEX), validando que los módulos de inventario, órdenes de producción, MRP, fichas, "
        "calidad, estadísticas, reportes, alertas y gestión de usuarios cumplan los requisitos "
        "funcionales (RF01–RF20) y no funcionales (RNF01–RNF08) definidos en la documentación del "
        "proyecto, mediante pruebas de interfaz de usuario (caja negra), pruebas de autorización por "
        "roles y verificación de persistencia de datos.",
    )
    add_para(
        doc,
        "Al finalizar las pruebas se debe disponer de un registro de casos ejecutados, no conformidades "
        "identificadas y un acta de recibido a satisfacción que habilite el paso a la fase de despliegue.",
    )

    # 2. DESCRIPCIÓN
    add_heading_doc(doc, "2. DESCRIPCIÓN DEL SISTEMA")
    add_para(
        doc,
        "SIPITEX es una aplicación web monolítica desarrollada en .NET 10 (ASP.NET Core MVC) con "
        "arquitectura por capas (Domain, Application, Infrastructure, Web), autenticación por cookies "
        "y persistencia en SQLite mediante Entity Framework Core. Está orientada a la operación del "
        "taller textil del centro de formación SENA CMTC (programa ADSO).",
    )
    add_para(
        doc,
        "El sistema permite gestionar materias primas e insumos, solicitudes de material a bodega, "
        "órdenes de producción, simulación MRP/BOM, registro de producción por fichas de aprendices, "
        "inspecciones de calidad, KPIs y gráficos, exportación de reportes PDF/Excel, alertas por "
        "correo electrónico y administración de usuarios con roles Administrador, Instructor y Bodeguero.",
    )
    add_para(
        doc,
        "Usuarios demo de prueba: admin@sipitex.test / Admin123! (Administrador); "
        "instructor@sipitex.test / Instructor123! (Instructor); bodega@sipitex.test / Bodega123! (Bodeguero).",
    )
    add_para(
        doc,
        "Fuera de alcance: facturación, nómina, integración con ERP externo y aplicación móvil nativa.",
    )

    # 3. MÓDULOS
    add_heading_doc(doc, "3. MÓDULOS DEL SISTEMA A PROBAR")
    add_para(doc, "Estructura del sistema a probar (módulos, submódulos y formularios):")
    add_table(
        doc,
        ["Módulo", "Submódulo / Sección", "Formularios / Acciones"],
        [
            ["Autenticación y cuenta", "Inicio de sesión", "Login (correo, contraseña)"],
            ["Autenticación y cuenta", "Recuperación de contraseña", "Olvidé contraseña; Nueva contraseña"],
            ["Autenticación y cuenta", "Perfil", "Mi perfil (datos, foto, cambio de clave)"],
            ["Autenticación y cuenta", "Usuarios (Admin)", "Listado; Crear usuario; Editar; Activar/Desactivar"],
            ["Inventario", "Materias primas", "Agregar material; Ajustar stock; Cambiar estado"],
            ["Inventario", "Solicitudes de material", "Crear solicitud; Aprobar; Rechazar"],
            ["Órdenes de producción", "Gestión de órdenes", "Crear orden + MRP; Avance +10u"],
            ["MRP / Materiales", "BOM y simulación", "Consultar BOM; Simular requerimiento neto"],
            ["Fichas & producción", "Fichas de aprendices", "Registrar ficha; Filtros; Sesión diaria; Registrar hoy"],
            ["Control de calidad", "Inspecciones", "Registrar inspección; Historial de resultados"],
            ["Estadísticas", "KPIs y gráficos", "Consulta de indicadores y Chart.js"],
            ["Reportes", "Exportaciones", "Inventario; Órdenes; Calidad; Dashboard KPI (PDF/Excel)"],
            ["Alertas", "Correo y preferencias", "Guardar preferencias; Evaluar y enviar; Historial"],
            ["Panel / Home", "Hub de módulos", "Panel SIPITEX; Privacy; Error"],
        ],
    )

    # 4. FORMULARIOS
    add_heading_doc(doc, "4. FORMULARIOS DEL APLICATIVO A PROBAR")
    add_table(
        doc,
        ["#", "Formulario / Pantalla", "Ruta", "Roles principales"],
        [
            ["F01", "Iniciar sesión", "/Account/Login", "Anónimo"],
            ["F02", "Olvidé mi contraseña", "/Account/ForgotPassword", "Anónimo"],
            ["F03", "Nueva contraseña", "/Account/ResetPassword", "Anónimo (token)"],
            ["F04", "Mi perfil", "/Account/Profile", "Autenticado"],
            ["F05", "Gestión de usuarios", "/Account/Users", "Administrador"],
            ["F06", "Crear usuario", "/Account/CreateUser", "Administrador"],
            ["F07", "Editar usuario", "/Account/EditUser", "Administrador"],
            ["F08", "Acceso denegado", "/Account/AccessDenied", "Autenticado"],
            ["F09", "Inventario de materias primas", "/Inventario", "Admin / Bodeguero / Instructor"],
            ["F10", "Órdenes de producción", "/Ordenes", "Admin / Bodeguero / Instructor"],
            ["F11", "MRP / Materiales", "/Mrp", "Admin / Bodeguero / Instructor"],
            ["F12", "Fichas & producción", "/Fichas", "Admin / Instructor"],
            ["F13", "Control de calidad", "/Calidad", "Admin / Instructor"],
            ["F14", "Estadísticas y KPIs", "/Estadisticas", "Autenticado"],
            ["F15", "Reportes", "/Reportes", "Autenticado"],
            ["F16", "Alertas por correo", "/Alertas", "Autenticado"],
            ["F17", "Panel SIPITEX", "/Home", "Autenticado"],
        ],
    )

    # 5. METODOLOGÍA
    add_heading_doc(doc, "5. METODOLOGÍA PARA LA APLICACIÓN DE LAS PRUEBAS")
    add_para(
        doc,
        "A continuación se detalla la metodología de trabajo a seguir durante el proceso de pruebas "
        "para el Sistema Integrado de Producción e Inventario Textil (SIPITEX).",
    )
    add_para(
        doc,
        "La metodología general a implementar es de pruebas de interfaz de usuario con enfoque de "
        "caja negra, complementada con verificación de políticas de autorización por rol y permisos "
        "extendidos (claims). Se verificará el correcto funcionamiento del sistema frente a los "
        "requerimientos iniciales documentados en docs/01-Requisitos.md y docs/08-IEEE830-Especificacion.md.",
    )
    add_para(doc, "Ambiente de pruebas:", bold=True)
    add_para(
        doc,
        "• Ejecución local: dotnet run en src/Sipitex.Web (HTTPS local) o Docker Compose en http://localhost:8080.\n"
        "• Base de datos SQLite de demostración (sipitex.db) con seed automático.\n"
        "• Pruebas automatizadas de apoyo en tests/Sipitex.Tests (dotnet test).\n"
        "• Registro manual de casos en el formato de plan de pruebas de usuario (Formato 1).",
    )
    add_para(doc, "Proceso:", bold=True)
    add_para(
        doc,
        "1. El grupo de pruebas ejecuta cada caso según el cronograma, con el responsable funcional "
        "del módulo (Administrador, Instructor o Bodeguero según aplique).\n"
        "2. Cada caso se diligencia en el Formato de Plan de Pruebas de Usuario (Formato 1), indicando "
        "módulo, submódulo, formulario, pasos, resultado esperado, resultado real y evidencia.\n"
        "3. Las no conformidades se reportan al equipo de desarrollo para corrección.\n"
        "4. Tras la entrega de una versión corregida, se realizan segundas pruebas hasta verificar "
        "que el 100% de las no conformidades fueron corregidas o aclaradas.\n"
        "5. Se genera un acta de recibido a satisfacción (Formato 2) firmada por los responsables.",
    )

    add_heading_doc(doc, "5.1 Formato de Plan de Pruebas de Usuario (Formato 1)", level=2)
    add_para(
        doc,
        "Plantilla a diligenciar por cada caso de prueba ejecutado:",
    )
    add_table(
        doc,
        ["Campo", "Descripción"],
        [
            ["Identificador CP", "CP-SIP-XXX (número del caso)"],
            ["Versión", "Versión del aplicativo bajo prueba"],
            ["Responsable", "Nombre del funcionario que realiza la prueba"],
            ["Nombre del caso de prueba", "Ej.: Aprobar solicitud de material con stock suficiente"],
            ["Módulo / Submódulo / Formulario", "Según secciones 3 y 4"],
            ["Descripción de la prueba", "Pasos detallados de ejecución"],
            ["Resultados esperados", "Comportamiento ideal según requisitos"],
            ["Resultados reales", "Lo observado en la ejecución"],
            ["Error", "Descripción del defecto si aplica"],
            ["Imagen", "Captura(s) de pantalla del error o evidencia"],
        ],
    )

    add_heading_doc(doc, "5.2 Resultados posibles de las pruebas", level=2)
    add_table(
        doc,
        ["Nombre del resultado", "Descripción", "Impacto"],
        [
            [
                "Correcta o superada",
                "El sistema funciona de acuerdo a lo solicitado en los requerimientos iniciales.",
                "—",
            ],
            [
                "Con no conformidades de diseño",
                "Inconsistencias de diseño del formulario y/o informes (ubicación de campos, tipografía, etc.).",
                "Medio",
            ],
            [
                "Con no conformidades de lógica",
                "El sistema no funciona como se especificó en los requerimientos iniciales.",
                "Alto",
            ],
            [
                "Requerimientos nuevos",
                "Requerimientos no solicitados al proveedor que afectan el funcionamiento.",
                "Alto",
            ],
        ],
    )

    add_para(
        doc,
        "Los usuarios funcionales (Administrador del taller, Instructores y Bodegueros) son los "
        "responsables directos de la evaluación y detección de errores, apoyados por el grupo de "
        "pruebas del proyecto ADSO.",
    )

    add_heading_doc(doc, "5.3 Aplicación de segundas pruebas", level=2)
    add_para(
        doc,
        "El equipo de desarrollo deberá entregar una versión con las correcciones a las no "
        "conformidades. Luego de desplegarla en el ambiente de pruebas (local o Docker), el grupo "
        "de pruebas junto con el responsable funcional verificará cada no conformidad. Finalmente "
        "se confirmará el correcto funcionamiento del módulo, submódulo o formulario, salidas "
        "(reportes), consultas y cálculos (MRP, KPIs, descuento de stock).",
    )
    add_para(
        doc,
        "Se generará un acta de recibido por módulo/evento con lo probado y aprobado (Formato 2), "
        "firmada por los responsables de las pruebas, el profesional funcional y el coordinador del "
        "proyecto, validando entrada de datos, procesamiento, almacenamiento y salidas del sistema.",
    )

    add_heading_doc(doc, "5.4 Casos de prueba funcionales prioritarios", level=2)
    add_table(
        doc,
        ["ID", "Caso", "Pasos resumidos", "Resultado esperado"],
        [
            ["T01", "Login administrador", "Ingresar admin@sipitex.test / Admin123!", "Acceso a Inventario; menú completo"],
            ["T02", "Agregar material", "Inventario → nombre + stock + unidad → Agregar", "Material aparece en tabla"],
            ["T03", "Alerta stock bajo", "Ajustar stock por debajo del mínimo", "Indicación visual / alerta de stock bajo"],
            ["T04", "Crear solicitud", "Orden + material + cantidad → Solicitar", "Estado Pendiente"],
            ["T05", "Aprobar solicitud", "Bodega/Admin aprueba con stock suficiente", "Stock descontado; estado Aprobada"],
            ["T06", "Rechazar solicitud", "Rechazar solicitud pendiente", "Estado Rechazada; stock sin cambio"],
            ["T07", "Crear orden", "Admin: producto, cantidad, fecha límite", "Nueva OP-xxx en EnProceso"],
            ["T08", "Producción +10u", "Botón +10u en orden (Admin/Instructor)", "Avance y consumo BOM"],
            ["T09", "Simular MRP", "Producto + cantidad → Calcular", "Líneas OK o déficit por material"],
            ["T10", "Registrar ficha", "Código, proceso, instructor, turno", "Ficha listada; Instructor auto-asociado"],
            ["T11", "Sesión de producción", "Registrar sesión o Registrar hoy", "Avance de orden actualizado"],
            ["T12", "Inspección calidad", "Orden + unidades + resultado", "Registro Aprobada/Reproceso/Rechazada"],
            ["T13", "KPIs", "Abrir Estadísticas", "Valores coherentes con datos"],
            ["T14", "Exportar reportes", "Reportes → PDF y Excel de cada tipo", "Archivos descargables válidos"],
            ["T15", "Preferencias alertas", "Guardar preferencias y Evaluar", "Envíos recientes registrados"],
            ["T16", "Crear usuario", "Admin → Usuarios → Crear Instructor/Bodeguero", "Usuario en listado; puede autenticarse"],
            ["T17", "Acceso denegado", "Bodeguero intenta /Fichas o /Calidad", "AccessDenied / sin menú"],
            ["T18", "Reset contraseña", "Olvidé contraseña → token → nueva clave", "Login exitoso con nueva clave"],
            ["T19", "Aprobar sin stock", "Aprobar solicitud con stock insuficiente", "Error controlado; sin descuento inválido"],
            ["T20", "Permisos extendidos", "Instructor con claim Solicitudes.Aprobar", "Puede aprobar según policy"],
        ],
    )

    # 6. CRONOGRAMA
    add_heading_doc(doc, "6. CRONOGRAMA DE IMPLEMENTACIÓN DE PRUEBAS")
    add_para(doc, "Tabla 1. Cronograma de implementación de pruebas.")
    add_table(
        doc,
        ["Semana", "Actividad", "Módulos", "Responsable"],
        [
            ["S1 — Día 1", "Preparación ambiente, smoke test y Login/Perfil", "Account, Home", "Grupo de pruebas"],
            ["S1 — Día 2", "Inventario y solicitudes", "Inventario", "Bodeguero + Grupo pruebas"],
            ["S1 — Día 3", "Órdenes y MRP", "Órdenes, MRP", "Admin + Grupo pruebas"],
            ["S1 — Día 4", "Fichas y Calidad", "Fichas, Calidad", "Instructor + Grupo pruebas"],
            ["S1 — Día 5", "Estadísticas, Reportes y Alertas", "Análisis", "Admin + Grupo pruebas"],
            ["S2 — Día 1", "Usuarios, roles y permisos", "Account/Users", "Administrador"],
            ["S2 — Día 2", "Casos negativos y seguridad", "Todos", "Grupo de pruebas"],
            ["S2 — Día 3", "Segundas pruebas / regresiones", "NC abiertas", "Grupo + Funcionales"],
            ["S2 — Día 4", "Acta de recibido y cierre", "—", "Coordinador + Funcionales"],
        ],
    )

    # 7. RESPONSABLES
    add_heading_doc(doc, "7. RESPONSABLES DE LAS PRUEBAS")
    add_table(
        doc,
        ["Rol en pruebas", "Perfil SIPITEX", "Alcance"],
        [
            ["Coordinador de pruebas", "Administrador / Líder ADSO", "Cronograma, actas, cierre"],
            ["Grupo de pruebas", "Equipo técnico ADSO", "Ejecución de casos, registro Formato 1"],
            ["Responsable funcional Bodega", "Bodeguero", "Inventario, solicitudes, stock"],
            ["Responsable funcional Producción", "Instructor", "Fichas, calidad, avance de órdenes"],
            ["Responsable funcional Sistema", "Administrador", "Usuarios, órdenes, alertas, reportes"],
            ["Desarrollo / Corrección NC", "Equipo de desarrollo", "Corregir no conformidades reportadas"],
        ],
    )

    # 8. RIESGOS
    add_heading_doc(doc, "8. RIESGOS")
    add_table(
        doc,
        ["Riesgo", "Impacto", "Mitigación"],
        [
            ["Ambiente de pruebas no disponible o DB corrupta", "Alto", "Usar Docker Compose y regenerar seed SQLite"],
            ["Indisponibilidad de responsables funcionales", "Alto", "Roles demo predefinidos; reasignar en cronograma"],
            ["SMTP no configurado para alertas reales", "Medio", "Validar outbox/demo y preferencias sin correo externo"],
            ["Diferencias entre docs (JWT) e implementación (cookies)", "Medio", "Probar según comportamiento real implementado"],
            ["No conformidades de lógica no corregidas a tiempo", "Alto", "Priorizar NC Alto; bloquear cierre hasta 100%"],
            ["Datos de prueba insuficientes para KPIs/reportes", "Bajo", "Ejecutar T01–T12 antes de Estadísticas/Reportes"],
        ],
    )

    # 9. RESPONSABILIDADES
    add_heading_doc(doc, "9. RESPONSABILIDADES")
    add_para(doc, "1. Matriz de responsabilidades")
    add_para(doc, "Tabla 2. Matriz de responsabilidades (R = Responsable, A = Apoya, C = Consulta).")
    add_table(
        doc,
        ["Responsabilidades", "Grupo de pruebas", "Admin funcional", "Instructor", "Bodeguero", "Desarrollo"],
        [
            ["1. Preparar ambiente y datos demo", "R", "A", "C", "C", "A"],
            ["2. Ejecutar casos de caja negra", "R", "A", "A", "A", "C"],
            ["3. Validar inventario y solicitudes", "A", "C", "A", "R", "C"],
            ["4. Validar fichas y calidad", "A", "C", "R", "—", "C"],
            ["5. Validar usuarios, órdenes y alertas", "A", "R", "C", "C", "C"],
            ["6. Registrar NC en Formato 1", "R", "A", "A", "A", "C"],
            ["7. Corregir no conformidades", "C", "C", "C", "C", "R"],
            ["8. Segundas pruebas y regresión", "R", "A", "A", "A", "A"],
            ["9. Firmar acta de recibido (Formato 2)", "A", "R", "R", "R", "C"],
        ],
    )

    add_heading_doc(doc, "10. CRITERIOS DE ACEPTACIÓN")
    add_para(
        doc,
        "• Compilación y ejecución sin errores críticos (dotnet build / docker compose).\n"
        "• Persistencia SQLite entre reinicios.\n"
        "• Casos T01–T20 ejecutados con resultado Correcta o NC de diseño menores aceptadas.\n"
        "• 100% de no conformidades de lógica corregidas o aclaradas formalmente.\n"
        "• Separación de capas mantenida (sin EF Core en controladores).\n"
        "• Acta de recibido firmada por responsables funcionales y coordinador.",
    )

    add_heading_doc(doc, "11. CONTROL DEL DOCUMENTO")
    add_table(
        doc,
        ["Versión", "Fecha", "Descripción", "Autor"],
        [
            ["1.0", "Julio 2026", "Plan de pruebas diligenciado para SIPITEX a partir de la plantilla institucional", "Equipo ADSO / SIPITEX"],
        ],
    )

    doc.save(path)
    print(f"DOCX: {path}")


def pdf_styles():
    styles = getSampleStyleSheet()
    styles.add(
        ParagraphStyle(
            name="CoverTitle",
            parent=styles["Title"],
            fontSize=16,
            textColor=colors.HexColor("#1A3A5C"),
            alignment=TA_CENTER,
            spaceAfter=12,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H1Sip",
            parent=styles["Heading1"],
            fontSize=13,
            textColor=colors.HexColor("#1A3A5C"),
            spaceBefore=14,
            spaceAfter=8,
        )
    )
    styles.add(
        ParagraphStyle(
            name="H2Sip",
            parent=styles["Heading2"],
            fontSize=11,
            textColor=colors.HexColor("#1A3A5C"),
            spaceBefore=10,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="BodySip",
            parent=styles["Normal"],
            fontSize=9.5,
            alignment=TA_JUSTIFY,
            leading=13,
            spaceAfter=6,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CenterSip",
            parent=styles["Normal"],
            fontSize=10,
            alignment=TA_CENTER,
            spaceAfter=4,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CellSip",
            parent=styles["Normal"],
            fontSize=7.5,
            leading=9,
            alignment=TA_LEFT,
        )
    )
    styles.add(
        ParagraphStyle(
            name="CellHdr",
            parent=styles["Normal"],
            fontSize=7.5,
            leading=9,
            textColor=colors.white,
            alignment=TA_CENTER,
        )
    )
    return styles


def pdf_table(headers, rows, styles, col_widths=None):
    data = [[Paragraph(h, styles["CellHdr"]) for h in headers]]
    for row in rows:
        data.append([Paragraph(str(c), styles["CellSip"]) for c in row])
    t = Table(data, colWidths=col_widths, repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1A3A5C")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#8899AA")),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 3),
                ("RIGHTPADDING", (0, 0), (-1, -1), 3),
                ("TOPPADDING", (0, 0), (-1, -1), 3),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F5F8")]),
            ]
        )
    )
    return t


def build_pdf(path: Path):
    styles = pdf_styles()
    doc = SimpleDocTemplate(
        str(path),
        pagesize=LETTER,
        leftMargin=1.8 * cm,
        rightMargin=1.8 * cm,
        topMargin=1.6 * cm,
        bottomMargin=1.6 * cm,
        title=TITLE,
        author="Equipo ADSO / SIPITEX",
    )
    story = []
    W = 17.4 * cm

    story.append(Paragraph("SENA — CMTC · Programa ADSO", styles["CenterSip"]))
    story.append(Paragraph("<b>SISTEMA INTEGRADO DE GESTIÓN — CONTROL DOCUMENTAL</b>", styles["CenterSip"]))
    story.append(Paragraph("GUÍA PARA EL DOCUMENTO DE PLAN DE PRUEBAS DE SOFTWARE", styles["CenterSip"]))
    story.append(Spacer(1, 8))
    story.append(Paragraph("Código: SIP-TIC-GUI-004 &nbsp;|&nbsp; Versión: 1.0 &nbsp;|&nbsp; Julio 2026", styles["CenterSip"]))
    story.append(Spacer(1, 20))
    story.append(Paragraph(TITLE, styles["CoverTitle"]))
    story.append(Paragraph(SUBTITLE, styles["CenterSip"]))
    story.append(Paragraph(
        "Documento diligenciado a partir de la plantilla institucional "
        "(SDS-TIC-GUI-004 / 114-GTI-GUI-08) y adaptado al sistema SIPITEX.",
        styles["BodySip"],
    ))
    story.append(PageBreak())

    story.append(Paragraph("1. OBJETIVO", styles["H1Sip"]))
    story.append(Paragraph(
        "Verificar el correcto funcionamiento del Sistema Integrado de Producción e Inventario Textil "
        "(SIPITEX), validando que los módulos de inventario, órdenes de producción, MRP, fichas, "
        "calidad, estadísticas, reportes, alertas y gestión de usuarios cumplan los requisitos "
        "funcionales (RF01–RF20) y no funcionales (RNF01–RNF08), mediante pruebas de interfaz de "
        "usuario (caja negra), autorización por roles y verificación de persistencia. Al finalizar "
        "se debe contar con registro de casos, no conformidades y acta de recibido a satisfacción.",
        styles["BodySip"],
    ))

    story.append(Paragraph("2. DESCRIPCIÓN DEL SISTEMA", styles["H1Sip"]))
    story.append(Paragraph(
        "SIPITEX es una aplicación web monolítica en .NET 10 (ASP.NET Core MVC) con arquitectura por "
        "capas (Domain, Application, Infrastructure, Web), autenticación por cookies y SQLite (EF Core). "
        "Gestiona el taller textil del SENA CMTC (ADSO): materias primas, solicitudes a bodega, órdenes "
        "de producción, MRP/BOM, fichas de aprendices, calidad, KPIs, reportes PDF/Excel, alertas y "
        "usuarios (Administrador, Instructor, Bodeguero). Usuarios demo: admin@sipitex.test / Admin123!; "
        "instructor@sipitex.test / Instructor123!; bodega@sipitex.test / Bodega123!. "
        "Fuera de alcance: facturación, nómina, ERP externo y app móvil.",
        styles["BodySip"],
    ))

    story.append(Paragraph("3. MÓDULOS DEL SISTEMA A PROBAR", styles["H1Sip"]))
    story.append(pdf_table(
        ["Módulo", "Submódulo", "Formularios / Acciones"],
        [
            ["Autenticación", "Login / Reset / Perfil", "Login; Olvidé clave; Nueva clave; Mi perfil"],
            ["Usuarios", "CRUD (Admin)", "Listado; Crear; Editar; Activar/Desactivar"],
            ["Inventario", "Materiales y solicitudes", "Agregar; Ajustar; Estado; Solicitar; Aprobar; Rechazar"],
            ["Órdenes", "Producción", "Crear orden + MRP; Avance +10u"],
            ["MRP", "BOM / Simulación", "Consultar BOM; Simular requerimiento neto"],
            ["Fichas", "Producción por ficha", "Registrar ficha; Sesión; Registrar hoy; Filtros"],
            ["Calidad", "Inspecciones", "Registrar inspección; Historial"],
            ["Estadísticas", "KPIs", "Indicadores y gráfico Chart.js"],
            ["Reportes", "Exportaciones", "Inventario; Órdenes; Calidad; Dashboard (PDF/Excel)"],
            ["Alertas", "Correo", "Preferencias; Evaluar; Historial de envíos"],
        ],
        styles,
        [3.2 * cm, 4.2 * cm, 10 * cm],
    ))

    story.append(Paragraph("4. FORMULARIOS DEL APLICATIVO A PROBAR", styles["H1Sip"]))
    story.append(pdf_table(
        ["#", "Formulario", "Ruta", "Roles"],
        [
            ["F01", "Iniciar sesión", "/Account/Login", "Anónimo"],
            ["F02", "Olvidé mi contraseña", "/Account/ForgotPassword", "Anónimo"],
            ["F03", "Nueva contraseña", "/Account/ResetPassword", "Anónimo (token)"],
            ["F04", "Mi perfil", "/Account/Profile", "Autenticado"],
            ["F05", "Gestión de usuarios", "/Account/Users", "Administrador"],
            ["F06", "Crear usuario", "/Account/CreateUser", "Administrador"],
            ["F07", "Editar usuario", "/Account/EditUser", "Administrador"],
            ["F08", "Inventario", "/Inventario", "Admin / Bodeguero / Instructor"],
            ["F09", "Órdenes de producción", "/Ordenes", "Admin / Bodeguero / Instructor"],
            ["F10", "MRP / Materiales", "/Mrp", "Admin / Bodeguero / Instructor"],
            ["F11", "Fichas & producción", "/Fichas", "Admin / Instructor"],
            ["F12", "Control de calidad", "/Calidad", "Admin / Instructor"],
            ["F13", "Estadísticas y KPIs", "/Estadisticas", "Autenticado"],
            ["F14", "Reportes", "/Reportes", "Autenticado"],
            ["F15", "Alertas por correo", "/Alertas", "Autenticado"],
            ["F16", "Panel SIPITEX", "/Home", "Autenticado"],
            ["F17", "Acceso denegado", "/Account/AccessDenied", "Autenticado"],
        ],
        styles,
        [1.2 * cm, 4.5 * cm, 5.5 * cm, 6.2 * cm],
    ))

    story.append(Paragraph("5. METODOLOGÍA PARA LA APLICACIÓN DE LAS PRUEBAS", styles["H1Sip"]))
    story.append(Paragraph(
        "Metodología de pruebas de interfaz de usuario con enfoque de caja negra, complementada con "
        "verificación de políticas de autorización por rol y permisos extendidos (claims). Ambiente: "
        "dotnet run en src/Sipitex.Web o Docker Compose (http://localhost:8080), base SQLite con seed, "
        "y suite automatizada en tests/Sipitex.Tests. Los casos se registran en el Formato 1; las no "
        "conformidades se envían a desarrollo; tras correcciones se ejecutan segundas pruebas hasta "
        "cerrar el 100% de NC y se firma el acta de recibido (Formato 2).",
        styles["BodySip"],
    ))

    story.append(Paragraph("5.1 Resultados posibles", styles["H2Sip"]))
    story.append(pdf_table(
        ["Resultado", "Descripción", "Impacto"],
        [
            ["Correcta o superada", "Funciona según requerimientos iniciales.", "—"],
            ["NC de diseño", "Inconsistencias de UI/informes (campos, tipografía, etc.).", "Medio"],
            ["NC de lógica", "No funciona como se especificó en los requerimientos.", "Alto"],
            ["Requerimientos nuevos", "Necesidades no solicitadas que afectan el sistema.", "Alto"],
        ],
        styles,
        [3.5 * cm, 11.4 * cm, 2.5 * cm],
    ))

    story.append(Paragraph("5.2 Casos de prueba funcionales prioritarios", styles["H2Sip"]))
    story.append(pdf_table(
        ["ID", "Caso", "Pasos", "Esperado"],
        [
            ["T01", "Login admin", "admin@sipitex.test / Admin123!", "Acceso Inventario; menú completo"],
            ["T02", "Agregar material", "Nombre + stock + unidad → Agregar", "Aparece en tabla"],
            ["T03", "Stock bajo", "Ajustar bajo el mínimo", "Alerta / indicación visual"],
            ["T04", "Crear solicitud", "Orden + material + cantidad", "Estado Pendiente"],
            ["T05", "Aprobar solicitud", "Aprobar con stock suficiente", "Stock descontado"],
            ["T06", "Rechazar solicitud", "Rechazar pendiente", "Rechazada; stock intacto"],
            ["T07", "Crear orden", "Producto, cantidad, fecha", "OP-xxx EnProceso"],
            ["T08", "Producción +10u", "Botón +10u", "Avance y consumo BOM"],
            ["T09", "Simular MRP", "Producto + cantidad", "OK o déficit"],
            ["T10", "Registrar ficha", "Código, proceso, turno", "Ficha listada"],
            ["T11", "Sesión producción", "Registrar hoy / sesión", "Avance de orden"],
            ["T12", "Inspección calidad", "Orden + resultado", "Registro en historial"],
            ["T13", "KPIs", "Abrir Estadísticas", "Valores coherentes"],
            ["T14", "Reportes", "Export PDF/Excel", "Archivos válidos"],
            ["T15", "Alertas", "Preferencias + Evaluar", "Envíos registrados"],
            ["T16", "Crear usuario", "Admin crea Instructor/Bodega", "Usuario usable"],
            ["T17", "Acceso denegado", "Bodeguero → Fichas/Calidad", "AccessDenied"],
            ["T18", "Reset clave", "Forgot → token → nueva", "Login con nueva clave"],
            ["T19", "Aprobar sin stock", "Aprobar con déficit", "Error controlado"],
            ["T20", "Permisos extendidos", "Instructor + claim Aprobar", "Policy permite acción"],
        ],
        styles,
        [1.3 * cm, 3.2 * cm, 5.5 * cm, 7.4 * cm],
    ))

    story.append(Paragraph("6. CRONOGRAMA DE IMPLEMENTACIÓN DE PRUEBAS", styles["H1Sip"]))
    story.append(Paragraph("Tabla 1. Cronograma de implementación de pruebas.", styles["BodySip"]))
    story.append(pdf_table(
        ["Semana", "Actividad", "Módulos", "Responsable"],
        [
            ["S1-D1", "Ambiente, smoke, Login/Perfil", "Account, Home", "Grupo de pruebas"],
            ["S1-D2", "Inventario y solicitudes", "Inventario", "Bodeguero + Grupo"],
            ["S1-D3", "Órdenes y MRP", "Órdenes, MRP", "Admin + Grupo"],
            ["S1-D4", "Fichas y Calidad", "Fichas, Calidad", "Instructor + Grupo"],
            ["S1-D5", "Estadísticas, Reportes, Alertas", "Análisis", "Admin + Grupo"],
            ["S2-D1", "Usuarios, roles y permisos", "Users", "Administrador"],
            ["S2-D2", "Casos negativos y seguridad", "Todos", "Grupo de pruebas"],
            ["S2-D3", "Segundas pruebas / regresiones", "NC abiertas", "Grupo + Funcionales"],
            ["S2-D4", "Acta de recibido y cierre", "—", "Coordinador + Funcionales"],
        ],
        styles,
        [2.2 * cm, 5.5 * cm, 4.2 * cm, 5.5 * cm],
    ))

    story.append(Paragraph("7. RESPONSABLES DE LAS PRUEBAS", styles["H1Sip"]))
    story.append(pdf_table(
        ["Rol en pruebas", "Perfil SIPITEX", "Alcance"],
        [
            ["Coordinador de pruebas", "Administrador / Líder ADSO", "Cronograma, actas, cierre"],
            ["Grupo de pruebas", "Equipo técnico ADSO", "Ejecución y Formato 1"],
            ["Funcional Bodega", "Bodeguero", "Inventario y solicitudes"],
            ["Funcional Producción", "Instructor", "Fichas, calidad, avance"],
            ["Funcional Sistema", "Administrador", "Usuarios, órdenes, alertas"],
            ["Desarrollo", "Equipo desarrollo", "Corrección de NC"],
        ],
        styles,
        [4.5 * cm, 5.5 * cm, 7.4 * cm],
    ))

    story.append(Paragraph("8. RIESGOS", styles["H1Sip"]))
    story.append(pdf_table(
        ["Riesgo", "Impacto", "Mitigación"],
        [
            ["Ambiente no disponible o DB corrupta", "Alto", "Docker Compose + regenerar seed"],
            ["Ausencia de responsables funcionales", "Alto", "Usuarios demo; reasignar cronograma"],
            ["SMTP no configurado", "Medio", "Validar outbox/demo sin correo externo"],
            ["Docs (JWT) vs implementación (cookies)", "Medio", "Probar comportamiento real"],
            ["NC de lógica sin corregir a tiempo", "Alto", "Priorizar Alto; bloquear cierre"],
            ["Datos insuficientes para KPIs", "Bajo", "Ejecutar T01–T12 antes de análisis"],
        ],
        styles,
        [7 * cm, 2.2 * cm, 8.2 * cm],
    ))

    story.append(Paragraph("9. RESPONSABILIDADES", styles["H1Sip"]))
    story.append(Paragraph("Tabla 2. Matriz de responsabilidades (R = Responsable, A = Apoya, C = Consulta).", styles["BodySip"]))
    story.append(pdf_table(
        ["Responsabilidad", "Grupo", "Admin", "Instructor", "Bodega", "Dev"],
        [
            ["1. Preparar ambiente y datos", "R", "A", "C", "C", "A"],
            ["2. Ejecutar casos caja negra", "R", "A", "A", "A", "C"],
            ["3. Validar inventario/solicitudes", "A", "C", "A", "R", "C"],
            ["4. Validar fichas y calidad", "A", "C", "R", "—", "C"],
            ["5. Validar usuarios/órdenes/alertas", "A", "R", "C", "C", "C"],
            ["6. Registrar NC (Formato 1)", "R", "A", "A", "A", "C"],
            ["7. Corregir no conformidades", "C", "C", "C", "C", "R"],
            ["8. Segundas pruebas / regresión", "R", "A", "A", "A", "A"],
            ["9. Firmar acta de recibido", "A", "R", "R", "R", "C"],
        ],
        styles,
        [6.4 * cm, 2.2 * cm, 2.2 * cm, 2.4 * cm, 2.2 * cm, 2 * cm],
    ))

    story.append(Paragraph("10. CRITERIOS DE ACEPTACIÓN", styles["H1Sip"]))
    story.append(Paragraph(
        "• Compilación y ejecución sin errores críticos.<br/>"
        "• Persistencia SQLite entre reinicios.<br/>"
        "• Casos T01–T20 ejecutados (Correcta o NC de diseño menores aceptadas).<br/>"
        "• 100% de NC de lógica corregidas o aclaradas.<br/>"
        "• Separación de capas (sin EF en controladores).<br/>"
        "• Acta de recibido firmada por funcionales y coordinador.",
        styles["BodySip"],
    ))

    story.append(Paragraph("11. CONTROL DEL DOCUMENTO", styles["H1Sip"]))
    story.append(pdf_table(
        ["Versión", "Fecha", "Descripción", "Autor"],
        [["1.0", "Julio 2026", "Plan de pruebas diligenciado para SIPITEX", "Equipo ADSO / SIPITEX"]],
        styles,
        [2.2 * cm, 2.8 * cm, 8.4 * cm, 4 * cm],
    ))

    doc.build(story)
    print(f"PDF: {path}")


def main():
    docx_name = "Plan_de_Pruebas_SIPITEX.docx"
    pdf_name = "Plan_de_Pruebas_SIPITEX.pdf"
    build_docx(OUT_DIR / docx_name)
    build_pdf(OUT_DIR / pdf_name)
    # Copias para descarga en artifacts
    import shutil

    shutil.copy2(OUT_DIR / docx_name, ARTIFACT_DIR / docx_name)
    shutil.copy2(OUT_DIR / pdf_name, ARTIFACT_DIR / pdf_name)
    print("Artifacts listos en", ARTIFACT_DIR)


if __name__ == "__main__":
    main()
